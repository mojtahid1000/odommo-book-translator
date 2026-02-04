"""
═══════════════════════════════════════════════════════════════
 অদম্য প্রেস — Book Translation Web App v3.0
 Features:
   - Per-page real-time progress bar
   - Password-protected API access
   - Compact DOCX (matches original book layout)
   - Multi-API: Anthropic + OpenAI + Gemini
   - Translator name mandatory with admin panel
═══════════════════════════════════════════════════════════════
"""

import streamlit as st
import fitz  # PyMuPDF
import re
import os
import io
import time
import hashlib
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ═══════════════════════════════════════════════════════════════
# PAGE CONFIG
# ═══════════════════════════════════════════════════════════════
st.set_page_config(page_title="অদম্য প্রেস — Book Translator", page_icon="📚", layout="wide")

# ═══════════════════════════════════════════════════════════════
# CUSTOM CSS
# ═══════════════════════════════════════════════════════════════
st.markdown("""
<style>
    .stApp { background-color: #0a0f0d; color: #e0e0e0; }
    .main-title { text-align: center; color: #4CAF50; font-size: 2.2rem; margin-bottom: 0; }
    .sub-title { text-align: center; color: #888; font-size: 1rem; margin-top: 0; }
    .stat-box { background: #1a2520; border: 1px solid #2d4a3e; border-radius: 10px; padding: 20px; text-align: center; }
    .stat-number { font-size: 2.2rem; font-weight: bold; color: #4CAF50; }
    .stat-label { font-size: 0.85rem; color: #888; }
    .cost-box { background: linear-gradient(135deg, #1a3a2a, #1a2520); border: 1px solid #4CAF50; border-radius: 10px; padding: 15px 20px; margin: 15px 0; color: #c0e0cc; }
    .translator-badge { background: linear-gradient(135deg, #1a2540, #1a2030); border: 1px solid #4a6fa5; border-radius: 10px; padding: 12px 20px; margin: 10px 0; color: #a0c0e0; }
    .translator-badge .name { font-weight: bold; color: #70b0ff; font-size: 1.05rem; }
    .progress-container { background: #1a2520; border: 1px solid #2d4a3e; border-radius: 12px; padding: 20px; margin: 15px 0; }
    .progress-header { display: flex; justify-content: space-between; align-items: center; margin-bottom: 10px; }
    .progress-header .left { font-size: 1rem; color: #e0e0e0; }
    .progress-header .right { font-size: 0.9rem; color: #4CAF50; font-weight: bold; }
    .progress-bar-outer { background: #0d1512; border-radius: 8px; height: 28px; overflow: hidden; }
    .progress-bar-inner { background: linear-gradient(90deg, #2d7a3e, #4CAF50, #66d470); height: 100%; border-radius: 8px; display: flex; align-items: center; justify-content: center; min-width: 40px; transition: width 0.3s ease; }
    .progress-bar-text { color: white; font-size: 0.8rem; font-weight: bold; text-shadow: 0 1px 2px rgba(0,0,0,0.5); }
    .progress-stats { display: flex; justify-content: space-between; margin-top: 10px; font-size: 0.8rem; color: #888; }
    .progress-stats .item { text-align: center; }
    .progress-stats .value { color: #4CAF50; font-weight: bold; font-size: 0.95rem; }
    .success-box { background: #1a3a2a; border: 1px solid #4CAF50; border-radius: 10px; padding: 20px; margin: 15px 0; }
    .lock-screen { text-align: center; padding: 60px 20px; }
    .lock-icon { font-size: 4rem; margin-bottom: 20px; }
    div[data-testid="stExpander"] { border: 1px solid #2d4a3e; border-radius: 10px; }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════
# SYSTEM PROMPT (shared across all APIs)
# ═══════════════════════════════════════════════════════════════
SYSTEM_PROMPT = """You are a professional English-to-Bangla book translator for অদম্য প্রেস (Odommo Press).

## TRANSLATION RULES

### Language Priority
- **Use English** for commonly understood terms: Focus, Energy, Goal, Priority, Distraction, Productivity, Mindset, Confidence, Resilience, Motivation, Discipline, Process, Comfort Zone, Emotion, Stress, Balance, Relationship, Communication, Trust, Challenge, Growth, Leadership, Strategy, Marketing, Brand, etc.
- **Use Bangla** for sentence structure, connectors, verbs, common everyday words, emotional language.
- **AVOID** complex Bangla: Use "Distraction" not "বিক্ষিপ্ততা", "Resilience" not "স্থিতিস্থাপকতা".

### CRITICAL FORMATTING
- **Bold**: **text** (double asterisks)
- *Italic*: *text* (single asterisks)
- Headings: # H1, ## H2, ### H3
- Numbered lists: ১. ২. ৩. (Bangla numerals)
- Bullets: • or -
- Quotes: > "text" — Author
- Chapter/section titles MUST be **bold**

### COMPACT OUTPUT — CRITICAL
- Do NOT add extra blank lines between paragraphs
- Do NOT add spacing that wasn't in the original
- Keep content DENSE and COMPACT — match the original book layout
- One English page = one translated section, no expansion
- Minimize whitespace. No decorative separators.

### OUTPUT FORMAT
=== পৃষ্ঠা [ORIGINAL PAGE NUMBER IN BANGLA] ===
[translated content — compact, no extra spacing]
---

IMPORTANT: Page number MUST match the ORIGINAL source PDF page number."""

# ═══════════════════════════════════════════════════════════════
# API PROVIDERS & MODELS
# ═══════════════════════════════════════════════════════════════
API_PROVIDERS = {
    "Anthropic (Claude)": {
        "models": {
            "Claude Sonnet 4.5 (Best)": "claude-sonnet-4-5-20250929",
            "Claude Haiku 4.5 (Fast)": "claude-haiku-4-5-20251001",
        },
        "key_prefix": "sk-ant-",
        "key_label": "Anthropic API Key",
        "key_help": "Get from console.anthropic.com",
    },
    "OpenAI (GPT)": {
        "models": {
            "GPT-4o (Best)": "gpt-4o",
            "GPT-4o Mini (Fast)": "gpt-4o-mini",
        },
        "key_prefix": "sk-",
        "key_label": "OpenAI API Key",
        "key_help": "Get from platform.openai.com",
    },
    "Google (Gemini)": {
        "models": {
            "Gemini 2.0 Flash": "gemini-2.0-flash",
            "Gemini 1.5 Pro": "gemini-1.5-pro",
        },
        "key_prefix": "AI",
        "key_label": "Google AI API Key",
        "key_help": "Get from aistudio.google.com",
    },
}

# Cost per token (input, output) per model
COST_MAP = {
    "claude-sonnet-4-5-20250929": (3.0, 15.0),
    "claude-haiku-4-5-20251001": (0.80, 4.0),
    "gpt-4o": (2.5, 10.0),
    "gpt-4o-mini": (0.15, 0.60),
    "gemini-2.0-flash": (0.10, 0.40),
    "gemini-1.5-pro": (1.25, 5.0),
}

# Per-page cost estimate
PAGE_COST_EST = {
    "claude-sonnet-4-5-20250929": 0.0114,
    "claude-haiku-4-5-20251001": 0.0035,
    "gpt-4o": 0.008,
    "gpt-4o-mini": 0.001,
    "gemini-2.0-flash": 0.0005,
    "gemini-1.5-pro": 0.005,
}


# ═══════════════════════════════════════════════════════════════
# API CALL FUNCTIONS
# ═══════════════════════════════════════════════════════════════

def call_anthropic(api_key, model, system, user_msg):
    """Call Anthropic Claude API."""
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    response = client.messages.create(
        model=model, max_tokens=4096, system=system,
        messages=[{"role": "user", "content": user_msg}]
    )
    text = response.content[0].text
    in_t = response.usage.input_tokens
    out_t = response.usage.output_tokens
    rates = COST_MAP.get(model, (3.0, 15.0))
    cost = (in_t * rates[0] / 1_000_000) + (out_t * rates[1] / 1_000_000)
    return text, in_t, out_t, cost


def call_openai(api_key, model, system, user_msg):
    """Call OpenAI GPT API."""
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model=model, max_tokens=4096,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user_msg}
        ]
    )
    text = response.choices[0].message.content
    in_t = response.usage.prompt_tokens
    out_t = response.usage.completion_tokens
    rates = COST_MAP.get(model, (2.5, 10.0))
    cost = (in_t * rates[0] / 1_000_000) + (out_t * rates[1] / 1_000_000)
    return text, in_t, out_t, cost


def call_gemini(api_key, model, system, user_msg):
    """Call Google Gemini API."""
    import google.generativeai as genai
    genai.configure(api_key=api_key)
    gmodel = genai.GenerativeModel(model, system_instruction=system)
    response = gmodel.generate_content(user_msg)
    text = response.text
    in_t = response.usage_metadata.prompt_token_count if hasattr(response, 'usage_metadata') else 0
    out_t = response.usage_metadata.candidates_token_count if hasattr(response, 'usage_metadata') else 0
    rates = COST_MAP.get(model, (0.10, 0.40))
    cost = (in_t * rates[0] / 1_000_000) + (out_t * rates[1] / 1_000_000)
    return text, in_t, out_t, cost


def translate_single_page(api_key, provider, model, page_num, page_text):
    """Translate a single page using the selected API provider."""
    user_msg = (
        f"Translate this page to Bangla. This is PAGE {page_num} — output as পৃষ্ঠা {int_to_bangla(page_num)}.\n"
        f"Keep ALL **bold**, *italic*, # heading formatting. Keep content COMPACT — no extra spacing.\n\n"
        f"--- PAGE {page_num} ---\n{page_text}"
    )

    if provider == "Anthropic (Claude)":
        return call_anthropic(api_key, model, SYSTEM_PROMPT, user_msg)
    elif provider == "OpenAI (GPT)":
        return call_openai(api_key, model, SYSTEM_PROMPT, user_msg)
    elif provider == "Google (Gemini)":
        return call_gemini(api_key, model, SYSTEM_PROMPT, user_msg)


# ═══════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════

def extract_pages(pdf_file, start_page, end_page):
    pdf_file.seek(0)
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    pages = []
    total = doc.page_count
    actual_end = min(end_page, total)
    for i in range(start_page - 1, actual_end):
        text = doc[i].get_text("text").strip()
        if text:
            pages.append((i + 1, text))
    doc.close()
    return pages, total


def parse_single_page(raw_text, expected_page_num):
    """Parse translation output for a single page."""
    pattern = r'===\s*পৃষ্ঠা\s*([০-৯]+)\s*==='
    parts = re.split(pattern, raw_text)
    if len(parts) > 1:
        page_num = parts[1]
        content = parts[2].strip() if len(parts) > 2 else ""
        content = re.sub(r'\n---\s*$', '', content).strip()
        return {"page": page_num, "content": content}
    else:
        # No marker found — use raw text with expected page num
        content = raw_text.strip()
        content = re.sub(r'\n---\s*$', '', content).strip()
        return {"page": int_to_bangla(expected_page_num), "content": content}


def bangla_to_int(s):
    m = {'০':'0','১':'1','২':'2','৩':'3','৪':'4','৫':'5','৬':'6','৭':'7','৮':'8','৯':'9'}
    try: return int("".join(m.get(c, c) for c in str(s)))
    except: return 0

def int_to_bangla(n):
    m = {'0':'০','1':'১','2':'২','3':'৩','4':'৪','5':'৫','6':'৬','7':'৭','8':'৮','9':'৯'}
    return "".join(m.get(c, c) for c in str(n))


def add_formatted_text(para, text, base_bold=False, base_italic=False, font_size=Pt(11), font_name='Noto Sans Bengali'):
    """Parse **bold**, *italic*, ***both*** and add runs."""
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)'
    last = 0
    for m in re.finditer(pattern, text):
        before = text[last:m.start()]
        if before:
            r = para.add_run(before); r.font.size = font_size; r.font.name = font_name
            r.bold = base_bold; r.italic = base_italic
        if m.group(2):
            r = para.add_run(m.group(2)); r.bold = True; r.italic = True
        elif m.group(3):
            r = para.add_run(m.group(3)); r.bold = True; r.italic = base_italic
        elif m.group(4):
            r = para.add_run(m.group(4)); r.bold = base_bold; r.italic = True
        r.font.size = font_size; r.font.name = font_name
        last = m.end()
    rem = text[last:]
    if rem:
        r = para.add_run(rem); r.font.size = font_size; r.font.name = font_name
        r.bold = base_bold; r.italic = base_italic


def build_docx(translated_pages, book_title, book_author, translator_name=""):
    """Build COMPACT DOCX matching original book layout."""
    doc = DocxDocument()

    style = doc.styles['Normal']
    style.font.name = 'Noto Sans Bengali'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # ── Title Page ──
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(book_title); r.font.size = Pt(24); r.bold = True; r.font.name = 'Noto Sans Bengali'

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(book_author); r.font.size = Pt(13); r.font.name = 'Noto Sans Bengali'

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("অদম্য প্রেস"); r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50); r.font.name = 'Noto Sans Bengali'

    if translator_name:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"অনুবাদক: {translator_name}"); r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66); r.font.name = 'Noto Sans Bengali'
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(datetime.now().strftime('%d %B %Y')); r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99); r.font.name = 'Noto Sans Bengali'

    doc.add_page_break()

    # ── Content Pages (COMPACT) ──
    for idx, page_data in enumerate(translated_pages):
        page_num = page_data["page"]
        content = page_data["content"]

        # Small page number — right aligned, minimal space
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"পৃষ্ঠা {page_num}"); r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA); r.italic = True; r.font.name = 'Noto Sans Bengali'

        lines = content.split('\n')
        skip_empty = False
        for line in lines:
            stripped = line.strip()

            # Skip consecutive empty lines (compact)
            if not stripped:
                if not skip_empty:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 0.5
                    skip_empty = True
                continue
            skip_empty = False

            if stripped.startswith('### '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
                add_formatted_text(p, stripped[4:], base_bold=True, font_size=Pt(11.5))
            elif stripped.startswith('## '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
                add_formatted_text(p, stripped[3:], base_bold=True, font_size=Pt(13))
            elif stripped.startswith('# '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
                add_formatted_text(p, stripped[2:], base_bold=True, font_size=Pt(14))
            elif stripped.startswith('> '):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                add_formatted_text(p, stripped[2:], base_italic=True, font_size=Pt(10.5))
            elif re.match(r'^[০-৯]+[\.\)]\s', stripped):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                add_formatted_text(p, stripped, font_size=Pt(10.5))
            elif stripped.startswith('• ') or stripped.startswith('- '):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                add_formatted_text(p, '• ' + stripped[2:], font_size=Pt(10.5))
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                add_formatted_text(p, stripped, font_size=Pt(10.5))

        # Page break between pages
        if idx < len(translated_pages) - 1:
            doc.add_page_break()

    # Footer
    if translator_name:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"অনুবাদ: {translator_name} | অদম্য প্রেস | {datetime.now().strftime('%Y')}")
        r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99); r.font.name = 'Noto Sans Bengali'

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════════════
# SESSION STATE
# ═══════════════════════════════════════════════════════════════
DEFAULTS = {
    "all_translated": [], "current_batch": 0, "translation_status": "idle",
    "logs": [], "total_cost": 0.0, "total_input_tokens": 0, "total_output_tokens": 0,
    "pages_data": [], "batch_result": [], "num_batches": 0, "total_pdf_pages": 0,
    "extract_hash": "", "authenticated": False, "page_progress": 0,
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = [] if isinstance(v, list) else v


# ═══════════════════════════════════════════════════════════════
# PASSWORD GATE
# ═══════════════════════════════════════════════════════════════
APP_PASSWORD = os.environ.get("APP_PASSWORD", "odommo2026")  # Set in Streamlit secrets or env

if not st.session_state.authenticated:
    st.markdown('<h1 class="main-title">📚 অদম্য প্রেস — Book Translator</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-title">English → Bangla | Powered by AI</p>', unsafe_allow_html=True)

    st.markdown("""
    <div class="lock-screen">
        <div class="lock-icon">🔐</div>
        <p style="color: #888; font-size: 1.1rem;">Enter password to access the translator</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        pwd = st.text_input("🔑 Password", type="password", placeholder="Enter access password...")
        if st.button("🔓 Unlock", type="primary", use_container_width=True):
            if pwd == APP_PASSWORD:
                st.session_state.authenticated = True
                st.rerun()
            else:
                st.error("❌ Wrong password. Contact admin for access.")

    st.markdown("<p style='text-align:center; color:#555; font-size:0.8rem; margin-top:40px;'>অদম্য প্রেস | Online Tech Academy</p>", unsafe_allow_html=True)
    st.stop()


# ═══════════════════════════════════════════════════════════════
# MAIN APP (after authentication)
# ═══════════════════════════════════════════════════════════════

st.markdown('<h1 class="main-title">📚 অদম্য প্রেস — Book Translator</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-title">English → Bangla | Anthropic • OpenAI • Gemini | Online Tech Academy</p>', unsafe_allow_html=True)
st.divider()

# ═══════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════
with st.sidebar:
    st.header("⚙️ Settings")

    # Translator name (mandatory)
    st.markdown("##### 👤 Translator (Required)")
    translator_name = st.text_input("Your Name", value="", placeholder="আপনার নাম...")

    st.divider()

    # API Provider selection
    st.markdown("##### 🤖 AI Provider")
    provider = st.selectbox("Provider", list(API_PROVIDERS.keys()))
    provider_info = API_PROVIDERS[provider]

    model_choice = st.selectbox("Model", list(provider_info["models"].keys()))
    model = provider_info["models"][model_choice]

    st.divider()

    # API Key
    env_key = os.environ.get("ANTHROPIC_API_KEY", "") if "Anthropic" in provider else ""
    if "OpenAI" in provider:
        env_key = os.environ.get("OPENAI_API_KEY", "")
    elif "Google" in provider:
        env_key = os.environ.get("GOOGLE_API_KEY", "")

    api_key = st.text_input(f"🔑 {provider_info['key_label']}", value=env_key,
                            type="password", help=provider_info['key_help'])

    st.divider()

    # Book metadata
    book_title = st.text_input("📖 Book Title (Bangla)", value="")
    book_author = st.text_input("✍️ Author", value="")

    st.divider()

    st.markdown("##### 📑 Page Range")
    c1, c2 = st.columns(2)
    with c1: start_page = st.number_input("Start", min_value=1, value=1)
    with c2: end_page = st.number_input("End", min_value=1, value=100)

    st.divider()
    batch_size = st.selectbox("📦 Review Every", [5, 10, 15, 20], index=1)

    st.divider()
    if st.button("🔄 Reset", use_container_width=True):
        for k, v in DEFAULTS.items():
            if k == "authenticated": continue
            st.session_state[k] = [] if isinstance(v, list) else v
        st.rerun()

    # Logout
    if st.button("🔒 Lock App", use_container_width=True, type="secondary"):
        st.session_state.authenticated = False
        st.rerun()

    st.divider()
    st.markdown(f"""**💰 Cost ({provider}):**\n\n| Model | /100 pages |\n|-------|:---:|\n""" +
                "\n".join(f"| {k} | ~${PAGE_COST_EST.get(v, 0.01)*100:.2f} |"
                          for k, v in provider_info["models"].items()))

# ═══════════════════════════════════════════════════════════════
# MAIN CONTENT
# ═══════════════════════════════════════════════════════════════

uploaded_file = st.file_uploader("📄 Upload English PDF", type=["pdf"])

if uploaded_file:
    # Re-extract on range change
    h = hashlib.md5(f"{uploaded_file.name}_{start_page}_{end_page}".encode()).hexdigest()
    if st.session_state.extract_hash != h:
        with st.spinner("📖 Extracting..."):
            pd, tp = extract_pages(uploaded_file, start_page, end_page)
            st.session_state.pages_data = pd
            st.session_state.total_pdf_pages = tp
            st.session_state.extract_hash = h
            if st.session_state.current_batch > 0:
                for k, v in DEFAULTS.items():
                    if k in ("authenticated", "extract_hash", "total_pdf_pages", "pages_data"): continue
                    st.session_state[k] = [] if isinstance(v, list) else v

    pages_data = st.session_state.pages_data
    num_pages = len(pages_data)
    num_batches = (num_pages + batch_size - 1) // batch_size
    st.session_state.num_batches = num_batches

    est_cost = num_pages * PAGE_COST_EST.get(model, 0.01)
    first_p = pages_data[0][0] if pages_data else start_page
    last_p = pages_data[-1][0] if pages_data else end_page

    # Translator badge
    if translator_name:
        st.markdown(f'<div class="translator-badge">👤 <span class="name">{translator_name}</span> '
                    f'<span style="color:#888">| {provider} — {model_choice} | Pages {first_p}–{last_p}</span></div>',
                    unsafe_allow_html=True)

    # Stats row
    cols = st.columns(5)
    stats = [
        ("📄 PDF Pages", st.session_state.total_pdf_pages),
        ("📑 To Translate", num_pages),
        (f"📍 Range", f"{first_p}–{last_p}"),
        ("📦 Batches", num_batches),
        ("💰 Est. Cost", f"${est_cost:.2f}"),
    ]
    for col, (label, val) in zip(cols, stats):
        sz = 'style="font-size:1.8rem;"' if "–" in str(val) else ""
        col.markdown(f'<div class="stat-box"><div class="stat-label">{label}</div>'
                     f'<div class="stat-number" {sz}>{val}</div></div>', unsafe_allow_html=True)

    st.markdown(f'<div class="cost-box">📊 <strong>{num_pages} pages</strong> (p{first_p}→{last_p}) '
                f'→ {num_batches} batches of {batch_size} | <strong>{model_choice}</strong> '
                f'| ⏸️ Review every {batch_size} pages</div>', unsafe_allow_html=True)

    # ─── State shortcuts ───
    status = st.session_state.translation_status
    current_batch = st.session_state.current_batch
    pages_done = len(st.session_state.all_translated)
    page_progress = st.session_state.page_progress

    # ─── Progress Bar (shows per-page progress) ───
    total_pages_overall = num_pages
    overall_done = pages_done + page_progress  # pages_done from prev batches + current batch progress
    pct = int((overall_done / total_pages_overall) * 100) if total_pages_overall > 0 else 0
    pct = min(pct, 100)

    if current_batch > 0 or page_progress > 0:
        last_done_page = first_p + overall_done - 1 if overall_done > 0 else first_p
        st.markdown(f"""
        <div class="progress-container">
            <div class="progress-header">
                <div class="left">📊 Translation Progress — {provider}</div>
                <div class="right">{pct}% Complete</div>
            </div>
            <div class="progress-bar-outer">
                <div class="progress-bar-inner" style="width: {max(pct, 2)}%;">
                    <span class="progress-bar-text">{pct}%</span>
                </div>
            </div>
            <div class="progress-stats">
                <div class="item">📦 Batches<br><span class="value">{current_batch}/{num_batches}</span></div>
                <div class="item">📄 Pages<br><span class="value">{overall_done}/{total_pages_overall}</span></div>
                <div class="item">📍 Done<br><span class="value">p{first_p}–{last_done_page}</span></div>
                <div class="item">⏳ Left<br><span class="value">{total_pages_overall - overall_done} pages</span></div>
                <div class="item">💰 Cost<br><span class="value">${st.session_state.total_cost:.4f}</span></div>
            </div>
        </div>
        """, unsafe_allow_html=True)

    # ─── START / CONTINUE ───
    if status in ["idle", "reviewing"]:
        if current_batch >= num_batches:
            st.session_state.translation_status = "complete"
            st.rerun()
        else:
            if not translator_name:
                st.warning("⚠️ **Enter your name** in the sidebar to start.")

            nsi = current_batch * batch_size
            nei = min(nsi + batch_size - 1, num_pages - 1)
            ns = pages_data[nsi][0]; ne = pages_data[nei][0]

            lbl = (f"🚀 Start — Batch 1/{num_batches} (p{ns}–{ne})" if status == "idle"
                   else f"▶️ Continue — Batch {current_batch+1}/{num_batches} (p{ns}–{ne})")

            if st.button(lbl, type="primary", use_container_width=True, disabled=(not translator_name)):
                if not api_key:
                    st.error(f"❌ Enter {provider_info['key_label']} in the sidebar.")
                else:
                    st.session_state.translation_status = "translating"
                    st.session_state.page_progress = 0
                    st.rerun()

    # ─── TRANSLATING (per-page with real-time progress) ───
    if status == "translating":
        batch_idx = st.session_state.current_batch
        b_start = batch_idx * batch_size
        b_end = min(b_start + batch_size, num_pages)
        batch_pages = pages_data[b_start:b_end]
        batch_count = len(batch_pages)

        st.markdown(f"### 🔄 Translating Batch {batch_idx+1}/{num_batches}")

        # Per-page progress
        progress_bar = st.progress(0)
        status_text = st.empty()
        page_results = []
        batch_cost = 0.0
        batch_in = 0
        batch_out = 0
        errors = []

        for i, (pg_num, pg_text) in enumerate(batch_pages):
            status_text.info(f"📝 Translating page {pg_num} ({i+1}/{batch_count})...")
            progress_bar.progress((i) / batch_count)

            try:
                raw, in_t, out_t, cost = translate_single_page(api_key, provider, model, pg_num, pg_text)
                parsed = parse_single_page(raw, pg_num)
                page_results.append(parsed)
                batch_cost += cost
                batch_in += in_t
                batch_out += out_t
                st.session_state.page_progress = i + 1
            except Exception as e:
                errors.append(f"Page {pg_num}: {str(e)}")
                page_results.append({"page": int_to_bangla(pg_num), "content": f"[Translation Error: {str(e)}]"})

            # Small delay to avoid rate limits
            if i < batch_count - 1:
                time.sleep(0.3)

        progress_bar.progress(1.0)
        status_text.success(f"✅ Batch {batch_idx+1} done — {len(page_results)} pages translated")

        # Store results
        st.session_state.all_translated.extend(page_results)
        st.session_state.batch_result = page_results
        st.session_state.total_cost += batch_cost
        st.session_state.total_input_tokens += batch_in
        st.session_state.total_output_tokens += batch_out
        st.session_state.current_batch += 1
        st.session_state.page_progress = 0

        page_nums = [p[0] for p in batch_pages]
        log = (f"✅ Batch {batch_idx+1}: p{page_nums[0]}–{page_nums[-1]} "
               f"— {len(page_results)} pages — ${batch_cost:.4f} "
               f"— {provider}/{model_choice} — by {translator_name} @ {datetime.now().strftime('%H:%M:%S')}")
        st.session_state.logs.append(log)
        if errors:
            for e in errors:
                st.session_state.logs.append(f"⚠️ {e}")

        st.session_state.translation_status = "reviewing"
        time.sleep(1)
        st.rerun()

    # ─── REVIEW MODE ───
    if status == "reviewing" and st.session_state.batch_result:
        bn = st.session_state.current_batch
        st.success(f"✅ **Batch {bn}/{num_batches} Complete** — Review, Download, or Continue.")

        st.markdown("### 📝 Review Translation")
        with st.expander(f"📖 Batch {bn} — Click to Review", expanded=True):
            for pd in st.session_state.batch_result:
                st.markdown(f"**━━━ পৃষ্ঠা {pd['page']} ━━━**")
                st.markdown(pd["content"])
                st.markdown("---")

        st.markdown("### 📥 Download DOCX")
        c1, c2 = st.columns(2)
        with c1:
            if st.session_state.all_translated:
                buf = build_docx(st.session_state.all_translated, book_title or "Book",
                                 book_author or "Author", translator_name)
                fp = st.session_state.all_translated[0]["page"]
                lp = st.session_state.all_translated[-1]["page"]
                st.download_button(f"📥 All ({len(st.session_state.all_translated)} pages: p{fp}–{lp})",
                                   data=buf, file_name=f"{book_title or 'book'}_p{bangla_to_int(fp)}-{bangla_to_int(lp)}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   use_container_width=True)
        with c2:
            if st.session_state.batch_result:
                buf2 = build_docx(st.session_state.batch_result, book_title or "Book",
                                  book_author or "Author", translator_name)
                bf = st.session_state.batch_result[0]["page"]
                bl = st.session_state.batch_result[-1]["page"]
                st.download_button(f"📥 Batch {bn} (p{bf}–{bl})",
                                   data=buf2, file_name=f"batch_{bn}_p{bangla_to_int(bf)}-{bangla_to_int(bl)}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   use_container_width=True)

    # ─── COMPLETE ───
    if status == "complete":
        st.balloons()
        st.markdown(f"""
        <div class="success-box">
            <h3 style="color:#4CAF50;margin-top:0;">🎉 Translation Complete!</h3>
            <p>📄 <strong>{len(st.session_state.all_translated)}</strong> pages | 📦 <strong>{st.session_state.current_batch}</strong> batches</p>
            <p>💰 <strong>${st.session_state.total_cost:.4f}</strong> | 🤖 {provider} — {model_choice}</p>
            <p>👤 <strong>{translator_name}</strong> | 🕐 {datetime.now().strftime('%d %b %Y, %I:%M %p')}</p>
        </div>
        """, unsafe_allow_html=True)

        if st.session_state.all_translated:
            buf = build_docx(st.session_state.all_translated, book_title or "Book",
                             book_author or "Author", translator_name)
            fp = st.session_state.all_translated[0]["page"]
            lp = st.session_state.all_translated[-1]["page"]
            st.download_button(f"📥 Download Complete (p{fp}–{lp})", data=buf, type="primary",
                               file_name=f"{book_title or 'book'}_complete.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               use_container_width=True)

    # ─── ADMIN PANEL ───
    if st.session_state.logs:
        with st.expander("📋 Admin Panel — Logs", expanded=False):
            st.markdown(f"""
            | Field | Value |
            |-------|-------|
            | 👤 Translator | **{translator_name or '—'}** |
            | 📖 Book | {book_title or '—'} |
            | 🤖 Provider | {provider} — {model_choice} (`{model}`) |
            | 📍 Range | Pages {start_page}–{end_page} |
            | 💰 Cost | ${st.session_state.total_cost:.4f} |
            | 🔤 Tokens | {st.session_state.total_input_tokens:,} in / {st.session_state.total_output_tokens:,} out |
            """)
            st.divider()
            for log in st.session_state.logs:
                if log.startswith("✅"): st.success(log)
                elif log.startswith("⚠️"): st.warning(log)
                else: st.error(log)

else:
    st.markdown("""
    ### 👋 How to Use

    1. **Enter password** to unlock the app
    2. **Enter your name** & **API key** in the sidebar
    3. **Select AI provider** — Anthropic, OpenAI, or Gemini
    4. **Upload** your English PDF
    5. **Set** page range and review batch size
    6. **Start** — watch per-page progress in real time
    7. **Review** each batch, **download** DOCX anytime
    8. **Continue** until complete

    ---

    ### ⭐ v3.0 Features

    | Feature | Details |
    |---------|---------|
    | 🔐 Password Protection | Only authorized users can access |
    | 📊 Real-time Progress | Bar moves for EACH page translated |
    | 📄 Compact DOCX | Matches original book layout — no extra spacing |
    | 🤖 Multi-API | Anthropic Claude, OpenAI GPT, Google Gemini |
    | ✅ Bold/Italic/Heading | All formatting preserved |
    | 👤 Translator Tracking | Name on DOCX cover & admin logs |

    ---
    *অদম্য প্রেস (Odommo Press) | Online Tech Academy*
    """)

st.divider()
st.markdown("<p style='text-align:center;color:#555;font-size:0.8rem;'>"
            "অদম্য প্রেস Book Translator v3.0 | Claude • GPT • Gemini | Online Tech Academy</p>",
            unsafe_allow_html=True)
